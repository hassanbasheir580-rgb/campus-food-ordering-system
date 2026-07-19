from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "input" / "Part_III_SEF_Final_template_2610.docx"
OUT = ROOT / "docs" / "generated"
DIAGRAMS = ROOT / "docs" / "diagrams"
SHOTS = ROOT / "docs" / "screenshots"

DOCX_OUT = OUT / "Part_III_System_Documentation.docx"
PDF_OUT = OUT / "Part_III_System_Documentation.pdf"

TEAM = [
    ("Hassan Abdelhaleem Hassan Basheir", "251UC2506T", "Student", "Student ordering, cart, payment, tracking, review"),
    ("Rashad Ali Qaid Sofan", "251UC2505G", "Vendor", "Menu and inventory, vendor queue, status updates, analytics"),
    ("Ahmad Aljammal", "251UC2500R", "Admin", "User/vendor/staff account management, configuration, reports"),
    ("Mohammed A. M. Alakklouk", "251UC250CR", "Staff", "Queue calling, pickup/delivery confirmation"),
]

DEMO_ACCOUNTS = [
    ("Student", "student@campus.test", "Student123!", "Browse menu, checkout, track order, review"),
    ("Vendor", "vendor@campus.test", "Vendor123!", "Manage queue, menu, inventory, analytics"),
    ("Admin", "admin@campus.test", "Admin123!", "Manage users, settings, reports"),
    ("Staff", "staff@campus.test", "Staff123!", "Call queue numbers, confirm pickup"),
]

IMAGES = {
    "use_case": DIAGRAMS / "use_case_diagram.png",
    "erd": DIAGRAMS / "erd_diagram.png",
    "state": DIAGRAMS / "order_state_diagram.png",
    "architecture": DIAGRAMS / "architecture_diagram.png",
    "deployment": DIAGRAMS / "deployment_diagram.png",
    "workflow": DIAGRAMS / "order_workflow_diagram.png",
    "login": SHOTS / "01-login-register.png",
    "student_menu": SHOTS / "02-student-menu.png",
    "cart": SHOTS / "03-cart-checkout.png",
    "tracking": SHOTS / "04-order-tracking.png",
    "vendor": SHOTS / "05-vendor-dashboard.png",
    "vendor_queue": SHOTS / "06-vendor-queue.png",
    "admin": SHOTS / "07-admin-dashboard.png",
    "staff": SHOTS / "08-staff-queue-panel.png",
}


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def clear_document(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def apply_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "Campus Food Ordering and Management System | System Documentation v3.0"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 112, 109)

    footer = section.footer.paragraphs[0]
    footer.text = "Part III SEF Final Submission"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 112, 109)


def add_para(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)
    return paragraph


def new_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing = [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))]
    num_id = (max(existing) if existing else 0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "0")
    num.append(abstract)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    p_pr.append(num_pr)


def add_bullets(doc: Document, items: Iterable[str]):
    num_id = new_numbering_instance(doc)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]):
    num_id = new_numbering_instance(doc)
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[Iterable[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_image(doc: Document, path: Path, caption: str, width: float = 6.35):
    if not path.exists():
        add_para(doc, f"[Image missing: {path.name}]")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(100, 112, 109)


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("System Documentation")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(31, 77, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("for")
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Campus Food Ordering and Management System")
    r.bold = True
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version 3.0")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(100, 112, 109)

    doc.add_paragraph()
    add_table(doc, ["Name", "Student No.", "Assigned Actor/Process"], [(m[0], m[1], m[2]) for m in TEAM], [2.6, 1.2, 2.7])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tutorial Section: TT5L | Group No.: 1 | Date: 16/06/2026")
    r.bold = True
    r.font.size = Pt(11)
    doc.add_page_break()


def add_front_matter(doc: Document):
    doc.add_heading("Contents", level=1)
    for item in [
        "1 Project Management",
        "2 System Overview",
        "3 Requirements",
        "4 Design",
        "5 Implementation Details",
        "6 Testing",
        "7 Sample Screens",
        "8 Reflection and Learning Outcomes",
        "9 User Guide",
        "10 Conclusion",
        "11 References",
    ]:
        add_para(doc, item)

    doc.add_heading("Revisions", level=1)
    add_table(
        doc,
        ["Version", "Primary Author(s)", "Description", "Date Completed"],
        [
            ("SRS v1.0", "All team members", "Part I requirements specification: actors, use cases, business rules, ERD, non-functional requirements.", "15/04/2026"),
            ("SDS v2.0", "All team members", "Part II design specification: activity, sequence, state, architecture, interface, component, and deployment design.", "02/06/2026"),
            ("System Documentation v3.0", "All team members", "Part III implementation documentation, prototype details, screenshots, testing, user guide, and reflections.", "16/06/2026"),
        ],
        [1.25, 1.65, 2.9, 1.0],
    )


def add_project_management(doc: Document):
    doc.add_heading("1 Project Management", level=1)
    doc.add_heading("1.1 Team Members", level=2)
    add_table(doc, ["Team Member", "Student No.", "Actor", "Assigned Processes"], TEAM, [2.1, 1.1, 1.0, 3.0])

    doc.add_heading("1.2 Problem Statement", level=2)
    add_para(
        doc,
        "Campus food outlets experience long queues and overcrowding during peak meal periods. Manual order handling, limited real-time inventory visibility, and lack of centralized pre-ordering reduce student break time and make vendor operations inefficient. The system solves this by digitizing menu browsing, ordering, scheduling, mock payment, queue tracking, vendor order management, staff pickup confirmation, and admin reporting in one role-based platform.",
    )

    doc.add_heading("1.3 Project Plan / Gantt Chart", level=2)
    add_table(
        doc,
        ["Activity", "Jun 2-4", "Jun 5-8", "Jun 9-11", "Jun 12-14", "Jun 15-16"],
        [
            ("Analyze Part I and Part II", "XXXX", "", "", "", ""),
            ("Backend API, database, seed data", "", "XXXX", "", "", ""),
            ("Frontend role dashboards", "", "XX", "XXXX", "", ""),
            ("Real-time queue, charts, validation", "", "", "XX", "XX", ""),
            ("Screenshots and diagrams", "", "", "", "XXXX", ""),
            ("Part III documentation and QA", "", "", "", "XX", "XXXX"),
        ],
        [2.1, 0.9, 0.9, 0.9, 0.9, 0.9],
    )


def add_system_overview(doc: Document):
    doc.add_heading("2 System Overview", level=1)
    doc.add_heading("2.1 Major Processes and Actors", level=2)
    add_table(
        doc,
        ["Actor", "Major Processes"],
        [
            ("Student", "Self-register or login with email/password, browse menu, add to cart, schedule order, make mock payment, track order/queue, submit review."),
            ("Vendor", "Login with an admin-created or seeded account, manage menu and inventory, accept/reject orders, update order status, view sales analytics."),
            ("Admin", "Create vendor/staff accounts, activate or suspend accounts, configure settings, view and export reports."),
            ("Staff", "Login with an admin-created or seeded account, call ready queue number, monitor queue display, confirm pickup/delivery, mark order completed."),
        ],
        [1.2, 5.2],
    )

    doc.add_heading("2.2 Tasks by Team Member", level=2)
    add_table(
        doc,
        ["Team Member", "Tasks Completed"],
        [
            ("Hassan", "Student workflow: menu browsing, cart, schedule, checkout/payment, tracking, review validation."),
            ("Rashad", "Vendor workflow: menu/inventory CRUD, queue management, status updates, vendor analytics."),
            ("Ahmad", "Admin workflow: users/vendors/staff table, account creation, status control, system settings, reports, CSV export, quality checks."),
            ("Mohammed", "Staff workflow: queue panel, queue calling, pickup/delivery confirmation, completion updates."),
        ],
        [1.4, 5.1],
    )

    doc.add_heading("2.3 Modules Developed", level=2)
    add_table(
        doc,
        ["Module", "Description"],
        [
            ("Authentication and RBAC", "Email/password login, database role resolution, student self-registration, admin-created vendor/staff accounts, JWT protected dashboards."),
            ("Menu and Inventory", "SQLite-backed vendor menu items, stock, availability, and time slot data."),
            ("Order and Payment", "Order creation, line items, configurable mock payment success/failure, payment receipts."),
            ("Queue Management", "Database queue tickets, public queue display, estimated waits, staff call and completion flow, live Socket.IO events."),
            ("Analytics and Reports", "Vendor revenue/top items/ratings and admin system-wide reports with CSV export."),
            ("Frontend UI", "Responsive React dashboards, charts, status chips, order stepper, toast messages, empty states."),
        ],
        [1.8, 4.8],
    )

    doc.add_heading("2.4 Assumptions Used in Part III", level=2)
    add_bullets(
        doc,
        [
            "The production payment gateway described in Part I/II is represented by a configurable mock payment gateway for the local prototype.",
            "SQLite is used for the local demo database; the SDS cloud deployment can replace it with a managed relational database later.",
            "Campus authentication and external notification services are represented by local JWT authentication and Socket.IO events.",
            "Only pickup fulfillment is fully demonstrated; the data model keeps a fulfillment type field for future delivery support.",
            "The admin account is seeded for the demo; vendors and staff do not self-register and are created by the campus administrator.",
        ],
    )

    doc.add_heading("2.5 Use Case Diagram", level=2)
    add_image(doc, IMAGES["use_case"], "Figure 1. Use case diagram based on Part I actors and Part II interface scope.", width=6.3)


def add_requirements(doc: Document):
    doc.add_heading("3 Requirements", level=1)
    doc.add_heading("3.1 Class Diagram / ERD", level=2)
    add_para(doc, "The ERD follows Part I business rules and Part II data design. User is the base authentication entity, while role-specific entities extend it through one-to-one profile tables.")
    add_image(doc, IMAGES["erd"], "Figure 2. ERD/class model for the implemented SQLite database.", width=6.3)

    doc.add_heading("3.2 State Diagram", level=2)
    add_para(doc, "The implemented order lifecycle follows the SDS state model: payment creates a confirmed or cancelled order; vendors move orders through preparation; staff call and complete queue tickets; reviews are only allowed after completion.")
    add_image(doc, IMAGES["state"], "Figure 3. Order lifecycle state diagram.", width=5.7)


def add_design(doc: Document):
    doc.add_heading("4 Design", level=1)
    doc.add_heading("4.1 Data Dictionary", level=2)
    add_table(
        doc,
        ["Entity", "Main Fields", "Purpose"],
        [
            ("User", "id, name, email, password_hash, phone, role, status", "Base authentication and authorization record."),
            ("Student", "user_id, student_no, faculty, department, dietary_preferences", "Student profile for orders and reviews."),
            ("Vendor", "id, user_id, outlet_name, operating_hours, avg_rating, verification_status, location", "Food outlet profile and owner of menu/order data."),
            ("Admin", "user_id, admin_level, department, access_permissions", "Administrative profile for system-wide operations."),
            ("Staff", "user_id, assigned_outlet, permissions", "Operational queue and pickup confirmation profile."),
            ("MenuItem", "id, vendor_id, name, description, price, category, prep_time, stock, is_available, image_url", "Vendor-managed sellable menu item."),
            ("Order", "id, student_id, vendor_id, status, fulfillment_type, total_price, scheduled_pickup_time", "Main transaction and status lifecycle record."),
            ("OrderItem", "id, order_id, menu_item_id, quantity, unit_price, subtotal, customizations", "Line items in an order."),
            ("Payment", "id, order_id, method, status, amount, transaction_ref, failure_reason", "Mock payment receipt and result."),
            ("QueueTicket", "id, order_id, queue_number, estimated_wait, status, called_at, served_at", "Queue display and staff operation record."),
            ("Review", "id, order_id, student_id, vendor_id, rating, comment, created_at", "Post-completion student feedback."),
            ("TimeSlot", "id, vendor_id, start_time, end_time, capacity, booked_count", "Scheduled pickup capacity."),
            ("SystemSetting", "key, value, type, description, updated_at", "Configurable runtime settings such as payment success rate and queue prefix."),
        ],
        [1.2, 2.8, 2.5],
    )

    doc.add_heading("4.2 Software Architecture", level=2)
    add_para(doc, "The prototype implements the SDS three-tier client-server architecture: React presentation tier, Express application tier, and SQLite data tier. Socket.IO provides real-time queue/order updates.")
    add_image(doc, IMAGES["architecture"], "Figure 4. Software architecture diagram.", width=6.3)

    doc.add_heading("4.3 Subsystems and Team Member Assignment", level=2)
    add_table(
        doc,
        ["Subsystem", "Team Member", "Implemented Scope"],
        [
            ("Student and Ordering", "Hassan", "Menu, cart, schedule, mock payment, tracking, review."),
            ("Vendor Management", "Rashad", "Menu/inventory, queue actions, status updates, sales analytics."),
            ("Admin Management", "Ahmad", "User/vendor control, settings, reports, export."),
            ("Staff and Queue", "Mohammed", "Queue display, call next, confirm pickup/delivery."),
        ],
        [1.8, 1.4, 3.3],
    )

    doc.add_heading("4.4 Main Screens", level=2)
    add_table(
        doc,
        ["Screen", "Actor", "Purpose"],
        [
            ("Login/Register", "All", "Email/password login, database role routing, student/customer self-registration, and live public queue display."),
            ("Student Menu", "Student", "Browse API-driven menu items with vendor/category filters."),
            ("Cart/Checkout", "Student", "Quantity control, schedule selection, payment method, order placement."),
            ("Order Tracking", "Student", "Status stepper, queue number, payment receipt, review form."),
            ("Vendor Dashboard", "Vendor", "Metrics, active queue, menu/inventory form and table, sales chart."),
            ("Admin Dashboard", "Admin", "System metrics, reports, create vendor/staff accounts, account status management, settings, export."),
            ("Staff Queue Panel", "Staff", "Large queue display, call ready tickets, complete called pickups."),
        ],
        [1.5, 1.0, 4.0],
    )

    doc.add_heading("4.5 Student Subsystem Screens", level=2)
    add_image(doc, IMAGES["student_menu"], "Figure 5. Student menu screen.", width=6.3)
    add_image(doc, IMAGES["cart"], "Figure 6. Student cart and checkout screen.", width=6.3)
    add_image(doc, IMAGES["tracking"], "Figure 7. Student order tracking screen.", width=6.3)

    doc.add_heading("4.6 Vendor/Admin/Staff Subsystem Screens", level=2)
    add_image(doc, IMAGES["vendor"], "Figure 8. Vendor dashboard and inventory screen.", width=6.3)
    add_image(doc, IMAGES["admin"], "Figure 9. Admin dashboard and reports screen.", width=6.3)
    add_image(doc, IMAGES["staff"], "Figure 10. Staff queue panel screen.", width=6.3)

    doc.add_heading("4.7 Main Components", level=2)
    add_table(
        doc,
        ["Component", "Subsystem", "Responsibility"],
        [
            ("AuthService", "User Management", "Email/password login, student self-registration, admin-created vendor/staff accounts, bcrypt password verification, JWT issuance."),
            ("OrderService", "Ordering", "Order creation, payment result handling, transitions, review rule enforcement."),
            ("MenuRepository", "Vendor Management", "Menu CRUD, availability, stock and time-slot operations."),
            ("QueueService/OrderRepository", "Queue", "Queue ticket creation, public queue display, call, served/cancelled state updates."),
            ("PublicQueueCard", "Frontend", "Displays the latest ready/called queue ticket from the public queue API and refreshes through Socket.IO/polling."),
            ("AnalyticsService", "Reporting", "Vendor and admin aggregate metrics and CSV export."),
            ("React Contexts", "Frontend", "Authenticated user session and cart state management."),
        ],
        [1.6, 1.4, 3.5],
    )

    doc.add_heading("4.8 Component Algorithm / Pseudocode / Workflow", level=2)
    add_para(doc, "Core order placement pseudocode:")
    add_para(
        doc,
        "1. Validate logged-in student and non-empty cart. 2. Load menu items from database. 3. Reject unavailable or insufficient-stock items. 4. Calculate total. 5. Run mock payment using SystemSetting.mockPaymentSuccessRate. 6. Create Order, OrderItem, and Payment records. 7. On success, decrement stock and create QueueTicket; on failure, save a failed payment and cancelled order without a queue ticket. 8. Emit Socket.IO events so student tracking, vendor queue, staff panel, and the public queue display refresh.",
    )
    add_image(doc, IMAGES["workflow"], "Figure 11. Main order workflow/activity diagram.", width=5.8)

    doc.add_heading("4.9 Deployment Diagram", level=2)
    add_image(doc, IMAGES["deployment"], "Figure 12. Local prototype and future cloud deployment diagram.", width=6.3)


def add_implementation(doc: Document):
    doc.add_heading("5 Implementation Details", level=1)
    doc.add_heading("5.1 Development Environment", level=2)
    add_table(
        doc,
        ["Area", "Tool / Library", "Reason Used"],
        [
            ("Frontend", "React, Vite, TypeScript", "Maintained SPA stack with fast local development and type safety."),
            ("Backend", "Node.js, Express, TypeScript", "REST API, middleware, validation, and service layering."),
            ("Database", "Node built-in SQLite module", "Real local SQLite persistence without native npm SQLite packages."),
            ("Real-time", "Socket.IO", "Order and queue status update broadcasts."),
            ("Charts", "Recharts", "Vendor/admin analytics charts in React."),
            ("Validation/Security", "Zod, bcryptjs, jsonwebtoken, helmet, cors", "Input validation, password hashing, JWT RBAC, HTTP hardening."),
            ("Visual QA", "Playwright", "Local browser screenshots for sample screens."),
            ("Documentation", "python-docx, ReportLab, Mermaid", "DOCX/PDF generation and diagram rendering."),
        ],
        [1.2, 2.0, 3.3],
    )

    doc.add_heading("5.2 Setup Instructions", level=2)
    add_numbered(
        doc,
        [
            "Install Node.js 24 or later.",
            "From the project root, run: npm install",
            "Initialize demo data: npm run seed",
            "Start both servers: npm run dev",
            "Open the frontend at http://localhost:5173",
            "Login with email/password only; the system reads the role from the database and redirects to the matching dashboard.",
            "Build and test: npm run build and npm run test",
        ],
    )

    doc.add_heading("5.3 Software Integration Strategy", level=2)
    add_para(doc, "Integration is API-driven. Frontend pages never contain hard-coded menu items, orders, analytics, users, or queue tickets. They call backend services through a typed API client. Login uses email/password only, and the backend determines the role from the users table. Role-based routes are protected by JWT and backend RBAC middleware. Socket.IO events and configured polling refresh student tracking, vendor queue, staff panels, and the public queue display when an order changes.")

    doc.add_heading("5.4 Database Implementation", level=2)
    add_para(doc, "The schema is created automatically by backend/src/database/schema.ts when the backend starts. The seed script populates the single demo admin, demo student, demo vendor and staff accounts, menu items, orders, payments, reviews, time slots, settings, and queue tickets. Student/customer self-registration creates only student-role accounts. Vendor and staff accounts are created by admin APIs and can be activated or suspended. Configurable runtime values are stored in system_settings.")
    add_table(doc, ["Role", "Email", "Password", "Purpose"], DEMO_ACCOUNTS, [1.0, 1.9, 1.3, 2.5])

    doc.add_heading("5.5 Important Files and Folders", level=2)
    add_table(
        doc,
        ["File / Folder", "Description"],
        [
            ("backend/src/config", "Environment configuration and runtime constants."),
            ("backend/src/constants", "Central role, order status, payment status, queue status, and transition enums."),
            ("backend/src/database", "SQLite schema, connection, and seed script."),
            ("backend/src/repositories", "Database access layer for users, menu, orders, settings."),
            ("backend/src/services", "Business rules for auth, orders, analytics, and real-time updates."),
            ("backend/src/routes", "REST API route definitions by domain and actor."),
            ("frontend/src/config", "App name, API URLs, nav labels, feature flags."),
            ("frontend/src/context", "Auth and cart state providers."),
            ("frontend/src/pages", "Role-based student, vendor, admin, and staff screens."),
            ("docs/diagrams", "Mermaid sources and rendered PNG diagrams."),
            ("docs/screenshots", "Playwright screenshots of the running application."),
            ("scripts", "Screenshot, Mermaid, and documentation generation scripts."),
        ],
        [2.0, 4.5],
    )


def add_testing(doc: Document):
    doc.add_heading("6 Testing", level=1)
    doc.add_heading("6.1 Testing Strategy", level=2)
    add_bullets(
        doc,
        [
            "Backend scenario test verifies health, email/password login for each role, restricted public registration, student checkout, vendor status updates, staff queue calling, review rules, admin account creation/suspension/password reset, role API blocking, and payment failure handling.",
            "Frontend TypeScript build verifies route/component/service integration.",
            "Full build verifies backend and frontend compile without errors.",
            "Playwright screenshot script verifies major pages render in a browser with real API data.",
            "Manual acceptance testing checks realistic role workflows, order transitions, public queue display updates, account management rules, and review rules.",
        ],
    )

    doc.add_heading("6.2 Test Data", level=2)
    add_table(doc, ["Role", "Email", "Password", "Seeded Data"], DEMO_ACCOUNTS, [1.0, 1.9, 1.3, 2.5])

    doc.add_heading("6.3 Acceptance Testing by Team Member / Module", level=2)
    add_table(
        doc,
        ["Team Member", "Module", "Acceptance Test", "Result"],
        [
            ("Hassan", "Student", "Login, browse menu, add item, view cart, choose schedule/payment, track active order, submit review only for completed order.", "Passed"),
            ("Rashad", "Vendor", "Login, view active queue, accept confirmed order, mark preparing order ready, update menu item, view analytics.", "Passed"),
            ("Ahmad", "Admin", "Login, create vendor account, create staff account, activate/suspend vendor and staff, protect seeded admin account, update configuration values, export CSV report.", "Passed"),
            ("Mohammed", "Staff", "Login, view ready queue ticket, call queue number, confirm called pickup and mark completed.", "Passed"),
            ("All", "Integration", "Order status updates are visible in student/vendor/staff dashboards and public queue display through API refresh and Socket.IO events.", "Passed"),
            ("All", "Build and Audit", "npm run build, npm run test, npm audit --audit-level=high.", "Passed"),
        ],
        [1.0, 1.2, 3.5, 0.8],
    )


def add_sample_screens(doc: Document):
    doc.add_heading("7 Sample Screens", level=1)
    for key, caption in [
        ("login", "Figure 13. Email/password login, student registration, and public queue display screen."),
        ("student_menu", "Figure 14. Student menu browsing screen."),
        ("cart", "Figure 15. Cart, schedule, and payment screen."),
        ("tracking", "Figure 16. Order tracking and queue status screen."),
        ("vendor_queue", "Figure 17. Vendor queue management screen."),
        ("admin", "Figure 18. Admin dashboard and analytics screen."),
        ("staff", "Figure 19. Staff queue panel screen."),
    ]:
        add_image(doc, IMAGES[key], caption, width=6.3)


def add_reflections(doc: Document):
    doc.add_heading("8 Reflection and Learning Outcomes", level=1)
    reflections = [
        ("Hassan", "Working on the Student subsystem showed how important a clear order flow is. A student should not need to understand backend states to place an order, so the UI uses menu filters, a simple cart, a payment result, and a stepper for tracking. I learned to connect user-facing convenience with database rules such as stock checking and review eligibility."),
        ("Rashad", "The Vendor subsystem required balancing speed and control. Vendors need quick accept, reject, ready, and inventory actions, but those actions must respect the order state machine. I learned how analytics, queue management, and inventory belong together because each order affects sales, stock, and queue visibility."),
        ("Ahmad", "The Admin subsystem made the value of centralized configuration clear. Settings such as mock payment success rate and queue prefix should not be hidden in components. I learned how reports, user control, and configuration support software quality assurance and operational visibility."),
        ("Mohammed", "The Staff subsystem focused on the last step of the service experience. Calling the correct queue number and confirming pickup must be simple, visible, and reliable. I learned that even a small staff panel needs strong status rules so completed orders are recorded accurately."),
    ]
    for name, text in reflections:
        doc.add_heading(name, level=2)
        add_para(doc, text)


def add_user_guide(doc: Document):
    doc.add_heading("9 User Guide", level=1)
    doc.add_heading("9.1 Login", level=2)
    add_numbered(doc, ["Open http://localhost:5173.", "Type the account email and password only; do not choose a role.", "Click Enter Dashboard.", "The backend validates the password, reads the role from the users table, and redirects to the Student, Vendor, Admin, or Staff dashboard.", "Use Register only for student/customer self-registration. Vendor and staff accounts are created by the admin dashboard."])

    doc.add_heading("9.2 Student Ordering", level=2)
    add_numbered(doc, ["Browse the Menu page and filter by vendor or category.", "Click Add on an available item.", "Open Cart.", "Adjust quantity if needed.", "Choose a pickup slot or leave as soon as possible.", "Select a payment method and place the order.", "Open Track to view status, queue number, payment receipt, and review eligibility."])

    doc.add_heading("9.3 Vendor Queue and Inventory", level=2)
    add_numbered(doc, ["Login as vendor.", "Review active orders in Live order queue.", "Accept confirmed orders to move them to Preparing.", "Mark preparing orders as Ready.", "Use the menu form and inventory table to add, edit, or remove menu items.", "Review sales analytics for revenue, order volume, top items, and ratings."])

    doc.add_heading("9.4 Admin Reports and Settings", level=2)
    add_numbered(doc, ["Login as admin.", "Review users, vendors, revenue, and order status charts.", "Use Create vendor to add an outlet name, email, phone, operating hours, and initial password.", "Use Create staff to add a staff name, email, phone, assigned outlet, and initial password.", "Activate or suspend vendor/staff accounts from the user management table; the seeded admin account is protected.", "Update settings such as mockPaymentSuccessRate and queuePrefix.", "Click Export CSV to download a report."])

    doc.add_heading("9.5 Staff Queue Handling", level=2)
    add_numbered(doc, ["Login as staff.", "Review Ready and Called tickets in the queue panel.", "Click Call next ready order or Call beside a ready ticket.", "Confirm pickup only after the queue number has been called.", "The order moves to Completed and the queue ticket is marked served.", "The public queue display and student tracking refresh from backend queue/order data."])


def add_conclusion(doc: Document):
    doc.add_heading("10 Conclusion", level=1)
    doc.add_heading("10.1 Completion Summary", level=2)
    add_para(doc, "Part III completed a working local prototype aligned with the Part I SRS and Part II SDS. The system includes email/password login with database role routing, student/customer self-registration, admin-created vendor/staff accounts, student ordering, mock payment, real queue tracking, vendor management, admin reporting, staff queue handling, seed data, diagrams, screenshots, and setup documentation.")

    doc.add_heading("10.2 Software Quality Assurance", level=2)
    add_bullets(doc, ["Central enums prevent duplicated status strings.", "Environment variables control API URLs, ports, app name, payment rate, queue prefix, and queue refresh interval.", "Seed data stores demo business records outside UI components.", "Repositories and services separate database access from business rules.", "RBAC middleware protects actor-specific API routes and account status is checked on every request.", "Build, scenario test, audit, and screenshot scripts support repeatable verification."])

    doc.add_heading("10.3 Collaboration", level=2)
    add_para(doc, "The final prototype follows the team actor assignments from Part I and subsystem assignments from Part II. Each role dashboard maps directly to the assigned responsibilities and acceptance tests.")

    doc.add_heading("10.4 Problems Encountered", level=2)
    add_bullets(doc, ["External services from the SRS, such as payment gateways and campus authentication, were not suitable for a local university prototype, so mock/local equivalents were documented.", "Real-time browser screenshots required local Playwright because the in-app browser surface was unavailable.", "LibreOffice was not available on PATH, so DOCX visual render QA could not be completed through the document renderer; a ReportLab PDF was generated separately."])

    doc.add_heading("10.5 Future Recommendations", level=2)
    add_bullets(doc, ["Replace mock payment with a sandbox payment gateway.", "Connect campus SSO and official notification services.", "Deploy with HTTPS, managed database backups, and production logging.", "Add automated end-to-end tests for every role workflow.", "Add delivery fulfillment and QR-code pickup verification."])

    doc.add_heading("11 References", level=1)
    add_bullets(
        doc,
        [
            "IEEE Std 830-1998: IEEE Recommended Practice for Software Requirements Specifications.",
            "Sommerville, I. (2016). Software Engineering (10th ed.). Pearson.",
            "Larman, C. (2004). Applying UML and Patterns (3rd ed.). Prentice Hall.",
            "OWASP Foundation (2024). Password Storage Cheat Sheet. https://cheatsheetseries.owasp.org",
            "World Wide Web Consortium (2018). Web Content Accessibility Guidelines (WCAG) 2.1.",
            "Payment Card Industry Security Standards Council (2022). PCI DSS v4.0.",
            "Amazon Web Services (2024). AWS Architecture Best Practices. https://aws.amazon.com/architecture",
            "React, Vite, TypeScript, Node.js, Express, SQLite, Socket.IO, Recharts, Mermaid, Playwright, python-docx, and ReportLab official documentation.",
            "Part I SRS: Proj_PI_TT5L_G1_Hassan, Rashad, Ahmad, Mohammed.pdf.",
            "Part II SDS: Proj_PII_TT5L_G1_Hassan, Rashad, Ahmad, Mohammed.pdf.",
        ],
    )


def generate_docx():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    apply_styles(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_project_management(doc)
    add_system_overview(doc)
    add_requirements(doc)
    add_design(doc)
    add_implementation(doc)
    add_testing(doc)
    add_sample_screens(doc)
    add_reflections(doc)
    add_user_guide(doc)
    add_conclusion(doc)
    doc.save(DOCX_OUT)


def md_files():
    implementation = """# Implementation Summary

The prototype implements the Campus Food Ordering and Management System from Part I and Part II.

## Stack
- Frontend: React + Vite + TypeScript
- Backend: Node.js + Express + TypeScript
- Database: SQLite through Node's maintained built-in `node:sqlite` module
- Real-time updates: Socket.IO
- Charts: Recharts
- Validation/security: Zod, bcryptjs, jsonwebtoken, helmet, cors

## Demo Accounts
| Role | Email | Password |
|---|---|---|
| Student | student@campus.test | Student123! |
| Vendor | vendor@campus.test | Vendor123! |
| Admin | admin@campus.test | Admin123! |
| Staff | staff@campus.test | Staff123! |

## Run
```bash
npm install
npm run seed
npm run dev
```

Frontend: http://localhost:5173
Backend health: http://localhost:4000/api/health

## Verification
```bash
npm run build
npm run test
npm audit --audit-level=high
node scripts/capture_screenshots.mjs
node scripts/render_mermaid.mjs
python scripts/build_part_iii_docs.py
```

## Notes
- Login is email/password based. The backend determines the role from the database and redirects the user to the correct dashboard.
- Only student/customer self-registration is public. Vendor and staff accounts are seeded or created by admin.
- The public queue display reads `/api/queue/display` and updates from real order/queue data through Socket.IO and polling.
- Mock payment success rate is configurable in the `system_settings` table and admin dashboard.
- Business data is seeded in `backend/src/database/seed.ts`, not hard-coded in frontend pages.
- UI pages load menu, orders, users, settings, queue tickets, and analytics through backend APIs.
"""

    test_plan = """# Test Plan

## Strategy
- Backend scenario tests: role login, restricted registration, student order/payment, vendor queue/status, staff call/complete, review rules, admin-created accounts, suspension/password reset, role API blocking, and payment failure.
- Frontend type check: `tsc -b`.
- Full build: backend TypeScript build and frontend Vite build.
- Visual verification: Playwright screenshots for all required screens.
- Manual acceptance: validate role workflows and status transitions.

## Acceptance Tests
| Module | Test | Expected Result |
|---|---|---|
| Student | Login, browse menu, add to cart, checkout, track, review | Student can complete the ordering workflow and review only completed orders. |
| Vendor | Accept/reject orders, mark ready, manage menu, view analytics | Vendor sees own outlet data and valid state transitions only. |
| Admin | Create vendor/staff, activate/suspend accounts, manage settings, reports, export CSV | Admin can control accounts, configuration, and reports while the seeded admin remains protected. |
| Staff | Call ready queue ticket and confirm called pickup | Queue number is called, public display updates, and only called orders can be completed. |
| Integration | Order updates across dashboards | Dashboards refresh through API and Socket.IO events. |

## Test Data
Seed data includes demo users, two vendors, seven menu items, sample orders, payments, queue tickets, reviews, time slots, and settings.
"""

    user_guide = """# User Guide

## Login
1. Start the app with `npm run dev`.
2. Open http://localhost:5173.
3. Type email and password only. Do not select a role.
4. The system reads the role from the database and redirects to the matching dashboard.
5. Register is for student/customer self-registration only.

## Student
1. Browse menu items.
2. Filter by vendor or category.
3. Add available items to cart.
4. Open Cart, choose pickup timing and payment method.
5. Place order.
6. Open Track to monitor status and queue number.
7. Submit a review after completion.

## Vendor
1. Open Vendor Dashboard.
2. Accept or reject confirmed orders.
3. Mark preparing orders as ready.
4. Add, edit, or delete menu items.
5. Review revenue, orders, top items, and ratings.

## Admin
1. Open Admin Dashboard.
2. Review system metrics and charts.
3. Create vendor accounts with outlet details and an initial password.
4. Create staff accounts with assigned outlet and an initial password.
5. Activate or suspend vendor/staff accounts.
6. Edit system settings.
7. Export CSV report.

## Staff
1. Open Staff Queue Panel.
2. Call a ready queue number.
3. Confirm pickup or delivery after the number has been called.
4. The order becomes completed.
"""

    (OUT / "implementation_summary.md").write_text(implementation, encoding="utf-8")
    (OUT / "test_plan.md").write_text(test_plan, encoding="utf-8")
    (OUT / "user_guide.md").write_text(user_guide, encoding="utf-8")


def pdf_image(path: Path, max_width: float = 6.5 * inch, max_height: float = 8.6 * inch):
    with PILImage.open(path) as image:
        width, height = image.size
    ratio = min(max_width / width, max_height / height, 1)
    return Image(str(path), width=width * ratio, height=height * ratio)


def generate_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverTitle", parent=styles["Title"], fontSize=24, leading=30, alignment=TA_CENTER, textColor=colors.HexColor("#1F4D78")))
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]
    body.spaceAfter = 6

    story = [
        Paragraph("System Documentation", styles["CoverTitle"]),
        Paragraph("Campus Food Ordering and Management System", styles["Title"]),
        Paragraph("Version 3.0 | Tutorial Section TT5L | Group 1 | 16/06/2026", styles["BodyText"]),
        Spacer(1, 0.25 * inch),
        Table([["Name", "Student No.", "Actor"]] + [[m[0], m[1], m[2]] for m in TEAM], colWidths=[3.1 * inch, 1.4 * inch, 1.5 * inch]),
        PageBreak(),
    ]

    for title, paragraphs in [
        ("1 Project Management", ["The project digitizes campus dining operations by supporting ordering, scheduling, mock payment, real-time queue tracking, admin-created vendor/staff accounts, vendor management, admin reporting, and staff pickup confirmation."]),
        ("2 System Overview", ["Actors are Student, Vendor, Admin, and Staff. Login uses email/password only; the backend determines role from the database. Each dashboard maps to the responsibilities defined in Part I and Part II."]),
        ("3 Requirements", ["The ERD and state model below summarize the implemented data and order lifecycle requirements."]),
    ]:
        story.append(Paragraph(title, h1))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, body))

    for path, caption in [
        (IMAGES["use_case"], "Use case diagram"),
        (IMAGES["erd"], "ERD/class diagram"),
        (IMAGES["state"], "Order state diagram"),
        (IMAGES["architecture"], "Software architecture"),
        (IMAGES["workflow"], "Order workflow"),
        (IMAGES["deployment"], "Deployment diagram"),
    ]:
        story.append(Paragraph(caption, h2))
        story.append(pdf_image(path))
        story.append(Spacer(1, 0.15 * inch))

    story.append(PageBreak())
    story.append(Paragraph("4 Design and Sample Screens", h1))
    for path, caption in [
        (IMAGES["login"], "Email/password login, student registration, and public queue display"),
        (IMAGES["student_menu"], "Student menu"),
        (IMAGES["cart"], "Cart and checkout"),
        (IMAGES["tracking"], "Order tracking"),
        (IMAGES["vendor"], "Vendor dashboard"),
        (IMAGES["admin"], "Admin dashboard"),
        (IMAGES["staff"], "Staff queue panel"),
    ]:
        story.append(Paragraph(caption, h2))
        story.append(pdf_image(path))
        story.append(Spacer(1, 0.15 * inch))

    story.append(PageBreak())
    story.append(Paragraph("5 Implementation, Testing, User Guide, and Conclusion", h1))
    for paragraph in [
        "The implementation uses React, Vite, TypeScript, Node.js, Express, SQLite, Socket.IO, Recharts, Zod, bcryptjs, JWT, Mermaid, Playwright, python-docx, and ReportLab.",
        "Setup: npm install, npm run seed, npm run dev. Verification: npm run build, npm run test, npm audit --audit-level=high.",
        "Demo accounts are student@campus.test, vendor@campus.test, admin@campus.test, and staff@campus.test with passwords documented in the DOCX and README. Vendor and staff accounts are created by the admin dashboard, not public registration.",
        "The public queue display reads real queue tickets from the backend and updates when vendors mark orders ready or staff call a queue number.",
        "The prototype is complete for local demonstration. Future work should integrate campus SSO, a sandbox payment gateway, production HTTPS deployment, and automated end-to-end tests.",
    ]:
        story.append(Paragraph(paragraph, body))

    SimpleDocTemplate(str(PDF_OUT), pagesize=letter, leftMargin=0.65 * inch, rightMargin=0.65 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch).build(story)


def main():
    generate_docx()
    md_files()
    generate_pdf()
    print(f"Generated {DOCX_OUT}")
    print(f"Generated {PDF_OUT}")
    print(f"Generated Markdown files in {OUT}")


if __name__ == "__main__":
    main()
